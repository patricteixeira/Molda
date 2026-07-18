from __future__ import annotations

import hashlib
from datetime import UTC, datetime
from pathlib import Path

import pytest
from docx import Document
from PIL import Image

from brand_runtime.ir.models import (
    BrandIR,
    BrandInfo,
    ColorToken,
    Evidence,
    FontToken,
    LogoAsset,
    RevisionInfo,
    SemanticRole,
)
from brand_runtime.roundtrip.docx import (
    DocxBrandError,
    analyze_docx_brand,
    apply_docx_brand_plan,
)


def _brand(tmp_path: Path) -> BrandIR:
    logo = tmp_path / "logo.png"
    Image.new("RGB", (180, 80), "#173F2C").save(logo)
    digest = hashlib.sha256(logo.read_bytes()).hexdigest()
    evidence = [
        Evidence(
            source_type="manual-entry",
            path="tests",
            confidence=1,
            authoritative=True,
            confirmed_at=datetime(2026, 7, 18, tzinfo=UTC),
        )
    ]
    return BrandIR(
        brand=BrandInfo(name="ACME"),
        revision=RevisionInfo(id="rev_docx", created_at=datetime(2026, 7, 18, tzinfo=UTC)),
        colors={
            "color.primary": ColorToken(value="#173F2C", evidence=evidence),
            "color.text": ColorToken(value="#1A1A1A", evidence=evidence),
        },
        fonts={
            "font.heading": FontToken(
                family="Aptos Display",
                weight=700,
                source="fallback",
                evidence=evidence,
            ),
            "font.body": FontToken(
                family="Aptos",
                weight=400,
                source="fallback",
                evidence=evidence,
            ),
        },
        roles={
            "heading": SemanticRole(
                font="font.heading",
                color="color.primary",
                min_size_px=24,
                max_size_px=48,
                line_height=1.1,
            ),
            "body": SemanticRole(
                font="font.body",
                color="color.text",
                min_size_px=14,
                max_size_px=20,
                line_height=1.5,
            ),
        },
        assets={
            "logo.primary": LogoAsset(
                path=str(logo),
                sha256=digest,
                format="png",
                evidence=evidence,
            )
        },
    )


def _source_docx(tmp_path: Path) -> Path:
    source = tmp_path / "proposta.docx"
    document = Document()
    document.add_paragraph("Proposta comercial")
    document.add_paragraph("Conteúdo que precisa continuar editável e intacto.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Etapa"
    table.cell(0, 1).text = "Prazo"
    table.cell(1, 0).text = "Entrega"
    table.cell(1, 1).text = "10 dias"
    document.sections[0].header.paragraphs[0].text = "Documento confidencial"
    document.save(source)
    return source


def test_docx_brand_plan_explains_changes_before_writing(tmp_path: Path):
    source = _source_docx(tmp_path)
    plan = analyze_docx_brand(source, _brand(tmp_path))

    assert plan.source.filename == "proposta.docx"
    assert plan.source.paragraph_count == 6
    assert plan.source.table_count == 1
    assert plan.brand_revision_id == "rev_docx"
    assert [operation.kind for operation in plan.operations] == [
        "document-styles",
        "paragraph-styles",
        "page-layout",
        "table-styles",
        "header-logo",
    ]


def test_docx_brand_apply_preserves_content_and_original(tmp_path: Path):
    source = _source_docx(tmp_path)
    original = source.read_bytes()
    brand = _brand(tmp_path)
    plan = analyze_docx_brand(source, brand)
    output = tmp_path / "proposta-molda.docx"

    result = apply_docx_brand_plan(source, output, plan, brand)

    assert source.read_bytes() == original
    assert result.content_preserved is True
    assert result.branded_sha256 == hashlib.sha256(output.read_bytes()).hexdigest()
    branded = Document(output)
    assert [paragraph.text for paragraph in branded.paragraphs] == [
        "Proposta comercial",
        "Conteúdo que precisa continuar editável e intacto.",
    ]
    assert branded.paragraphs[0].style.name == "Molda Título"
    assert branded.tables[0].cell(1, 0).text == "Entrega"
    assert branded.sections[0].header.paragraphs[0].text == "Documento confidencial"
    assert branded.sections[0].header._element.xpath(".//w:drawing")


def test_docx_brand_plan_is_bound_to_exact_source_bytes(tmp_path: Path):
    source = _source_docx(tmp_path)
    brand = _brand(tmp_path)
    plan = analyze_docx_brand(source, brand)
    changed = Document(source)
    changed.add_paragraph("Mudança posterior")
    changed.save(source)

    with pytest.raises(DocxBrandError, match="bytes do Word mudaram"):
        apply_docx_brand_plan(source, tmp_path / "out.docx", plan, brand)
