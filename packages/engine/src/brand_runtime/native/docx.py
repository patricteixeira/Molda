"""Renderer DOCX template-fill com estilos semânticos da marca."""

from __future__ import annotations

import os
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from brand_runtime.ir.models import BrandIR
from brand_runtime.kit.models import ContentSpec, LayoutSpec, TextValue
from brand_runtime.native.ooxml import OoxmlError, validate_ooxml


class DocxRenderError(OoxmlError):
    """O template ou os contratos não permitem um DOCX semântico coerente."""


def _style_name(role: str) -> str:
    words = role.replace("_", "-").split("-")
    return "Brand " + " ".join(word.capitalize() for word in words)


def _ensure_style(document: Document, role_name: str, ir: BrandIR):
    if role_name not in ir.roles:
        raise DocxRenderError(f"O papel semântico «{role_name}» não existe no Brand IR.")
    semantic_role = ir.roles[role_name]
    font_token = ir.fonts.get(semantic_role.font)
    color_token = ir.colors.get(semantic_role.color)
    if font_token is None or color_token is None:
        raise DocxRenderError(f"O papel semântico «{role_name}» possui referência ausente.")
    name = _style_name(role_name)
    try:
        style = document.styles[name]
    except KeyError:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = font_token.family
    style.font.size = Pt(semantic_role.max_size_px * 0.75)
    style.font.bold = font_token.weight >= 600
    style.font.italic = font_token.style == "italic"
    style.font.color.rgb = RGBColor.from_string(color_token.value.removeprefix("#"))
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_token.family)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font_token.family)
    return style


def _iter_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _replace_text_slot(
    document: Document, slot_id: str, value: TextValue, role: str, ir: BrandIR
) -> bool:
    token = f"{{{{slot:{slot_id}}}}}"
    for paragraph in _iter_paragraphs(document):
        if paragraph.text.strip() != token:
            continue
        paragraph.clear()
        paragraph.add_run(value.text)
        paragraph.style = _ensure_style(document, role, ir)
        return True
    return False


def _replace_logo_slot(document: Document, slot_id: str, asset: Path) -> bool:
    token = f"{{{{slot:{slot_id}}}}}"
    for paragraph in _iter_paragraphs(document):
        if paragraph.text.strip() != token:
            continue
        paragraph.clear()
        paragraph.add_run().add_picture(str(asset), width=Inches(1.15))
        return True
    return False


def _resolve_asset(path: str, asset_root: Path | None) -> Path:
    candidate = Path(path)
    if not candidate.is_absolute() and asset_root is not None:
        candidate = asset_root / candidate
    candidate = candidate.resolve()
    if not candidate.is_file():
        raise DocxRenderError(f"O asset «{path}» não foi encontrado.")
    return candidate


def render_docx(
    template_path: Path,
    output_path: Path,
    ir: BrandIR,
    layout: LayoutSpec,
    content: ContentSpec,
    *,
    asset_root: Path | None = None,
) -> Path:
    """Preenche placeholders nomeados e mantém texto, estilos e imagens nativos."""
    template_path = template_path.resolve()
    output_path = output_path.resolve()
    if template_path == output_path:
        raise DocxRenderError("O template original nunca pode ser sobrescrito.")
    blocking = [item for item in validate_ooxml(template_path) if item.blocking]
    if blocking:
        raise DocxRenderError(f"O template possui {len(blocking)} erro(s) estrutural(is).")
    if content.brand_revision_id != ir.revision.id:
        raise DocxRenderError("O Content Spec não pertence à revisão de marca informada.")
    if content.layout_id != layout.id:
        raise DocxRenderError("O Content Spec não pertence ao Layout Spec informado.")

    document = Document(template_path)
    missing: list[str] = []
    for slot in layout.slots:
        value = content.values.get(slot.id)
        if slot.kind == "text" and isinstance(value, TextValue):
            if not _replace_text_slot(document, slot.id, value, slot.role, ir):
                missing.append(slot.id)
        elif slot.kind == "logo":
            asset_token = slot.asset_token or next(iter(ir.assets), None)
            if asset_token is None or asset_token not in ir.assets:
                raise DocxRenderError("O layout exige logo, mas o Brand IR não possui asset.")
            asset = _resolve_asset(ir.assets[asset_token].path, asset_root)
            if not _replace_logo_slot(document, slot.id, asset):
                missing.append(slot.id)
    if missing:
        slots = ", ".join(f"«{slot_id}»" for slot_id in missing)
        raise DocxRenderError(f"O template não contém os placeholders obrigatórios {slots}.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    handle, temp_name = tempfile.mkstemp(
        prefix=f".{output_path.stem}-", suffix=".docx", dir=output_path.parent
    )
    os.close(handle)
    temp_path = Path(temp_name)
    try:
        document.save(temp_path)
        blocking = [item for item in validate_ooxml(temp_path) if item.blocking]
        if blocking:
            raise DocxRenderError(f"O DOCX gerado possui {len(blocking)} erro(s) estrutural(is).")
        os.replace(temp_path, output_path)
    finally:
        temp_path.unlink(missing_ok=True)
    return output_path
