"""Análise e aplicação não destrutiva de marca em documentos DOCX existentes."""

from __future__ import annotations

import hashlib
import os
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Literal

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pydantic import Field

from brand_runtime.ir.models import BrandIR, CamelModel, SemanticRole
from brand_runtime.native.ooxml import OoxmlError, validate_ooxml


class DocxBrandError(OoxmlError):
    """O DOCX ou o plano não permitem uma aplicação de marca segura."""


class DocxBrandSource(CamelModel):
    """Identidade e contagens imutáveis do Word analisado."""

    filename: str
    sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    size_bytes: int = Field(gt=0)
    paragraph_count: int = Field(ge=0)
    table_count: int = Field(ge=0)
    section_count: int = Field(ge=1)


class DocxBrandOperation(CamelModel):
    """Mudança agregada comunicada antes de qualquer escrita no arquivo."""

    id: str
    kind: Literal[
        "document-styles",
        "paragraph-styles",
        "page-layout",
        "table-styles",
        "header-logo",
    ]
    label_pt: str
    affected_count: int = Field(ge=0)
    target_role: str | None = None


class DocxBrandPlan(CamelModel):
    """Plano reproduzível vinculado aos bytes e à revisão de marca."""

    schema_version: Literal["0.1.0"] = "0.1.0"
    source: DocxBrandSource
    brand_revision_id: str
    operations: list[DocxBrandOperation]
    warnings: list[str]


class DocxBrandResult(CamelModel):
    """Evidência de aplicação, preservação textual e publicação do novo DOCX."""

    schema_version: Literal["0.1.0"] = "0.1.0"
    source_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    branded_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    output_filename: str
    applied_operation_ids: list[str]
    content_preserved: bool
    content_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")


@dataclass(frozen=True, slots=True)
class _ParagraphAssignment:
    """Classificação interna que preserva a estrutura nativa do parágrafo."""

    paragraph: object
    role_name: str
    force_custom_style: bool = False


_CUSTOM_HEADING_STYLE = "Molda Título"
_CUSTOM_BODY_STYLE = "Molda Texto"
_FONT_XML_TAGS = ("w:rFonts", "w:sz", "w:szCs", "w:color")


def _sha256(path: Path) -> str:
    with path.open("rb") as source_file:
        return hashlib.file_digest(source_file, "sha256").hexdigest()


def _all_body_paragraphs(document: Document) -> list:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def _content_digest(document: Document) -> str:
    """Resume texto de corpo, tabelas, cabeçalhos e rodapés sem serialização OOXML."""
    parts = [f"body:{paragraph.text}" for paragraph in document.paragraphs if paragraph.text]
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        parts.append(
                            f"table:{table_index}:{row_index}:{cell_index}:{paragraph.text}"
                        )
    seen_parts: set[str] = set()
    for section_index, section in enumerate(document.sections):
        for region_name, region in (("header", section.header), ("footer", section.footer)):
            part_name = str(region.part.partname)
            if part_name in seen_parts:
                continue
            seen_parts.add(part_name)
            for paragraph in region.paragraphs:
                if paragraph.text:
                    parts.append(f"{region_name}:{section_index}:{paragraph.text}")
    payload = "\n".join(parts).encode("utf-8")
    return hashlib.sha256(payload).hexdigest()


def _media_hashes(path: Path) -> set[str]:
    try:
        with zipfile.ZipFile(path) as package:
            return {
                hashlib.sha256(package.read(name)).hexdigest()
                for name in package.namelist()
                if name.startswith("word/media/") and not name.endswith("/")
            }
    except (OSError, zipfile.BadZipFile) as error:  # pragma: no cover - validado antes
        raise DocxBrandError("Não foi possível conferir as imagens do DOCX.") from error


def _semantic_role(ir: BrandIR, preferred: str) -> tuple[str, SemanticRole]:
    if preferred in ir.roles:
        return preferred, ir.roles[preferred]
    if not ir.roles:
        raise DocxBrandError("A revisão da marca não possui papéis tipográficos.")
    name = sorted(ir.roles)[0]
    return name, ir.roles[name]


def _looks_like_plain_title(text: str) -> bool:
    normalized = text.strip()
    return bool(normalized) and len(normalized) <= 120 and normalized[-1:] not in ".,;:!?"


def _assignments(document: Document, ir: BrandIR) -> list[_ParagraphAssignment]:
    heading_name, _ = _semantic_role(ir, "heading")
    body_name, _ = _semantic_role(ir, "body")
    paragraphs = _all_body_paragraphs(document)
    first_non_empty = next((item for item in paragraphs if item.text.strip()), None)
    assignments: list[_ParagraphAssignment] = []
    for paragraph in paragraphs:
        if not paragraph.text.strip():
            continue
        style_name = paragraph.style.name if paragraph.style is not None else "Normal"
        normalized_style = style_name.casefold()
        explicit_heading = normalized_style == "title" or normalized_style.startswith("heading ")
        inferred_heading = (
            paragraph is first_non_empty
            and normalized_style in {"normal", "body text"}
            and _looks_like_plain_title(paragraph.text)
        )
        assignments.append(
            _ParagraphAssignment(
                paragraph=paragraph,
                role_name=heading_name if explicit_heading or inferred_heading else body_name,
                force_custom_style=inferred_heading,
            )
        )
    return assignments


def _target_size_pt(role_name: str, role: SemanticRole) -> float:
    preferred_px = 36 if role_name == "heading" else 16
    bounded_px = min(role.max_size_px, max(role.min_size_px, preferred_px))
    return bounded_px * 0.75


def _configure_style(style, role_name: str, ir: BrandIR) -> None:
    role = ir.roles[role_name]
    font = ir.fonts.get(role.font)
    color = ir.colors.get(role.color)
    if font is None or color is None:
        raise DocxBrandError(f"O papel «{role_name}» referencia tokens ausentes.")
    style.font.name = font.family
    style.font.size = Pt(_target_size_pt(role_name, role))
    style.font.bold = font.weight >= 600
    style.font.italic = font.style == "italic"
    style.font.color.rgb = RGBColor.from_string(color.value.removeprefix("#"))
    run_properties = style._element.get_or_add_rPr()
    run_properties.rFonts.set(qn("w:ascii"), font.family)
    run_properties.rFonts.set(qn("w:hAnsi"), font.family)
    run_properties.rFonts.set(qn("w:eastAsia"), font.family)
    style.paragraph_format.line_spacing = role.line_height
    style.paragraph_format.space_after = Pt(10 if role_name == "heading" else 6)
    style.paragraph_format.keep_with_next = role_name == "heading"


def _ensure_custom_style(document: Document, name: str, role_name: str, ir: BrandIR):
    try:
        style = document.styles[name]
    except KeyError:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.hidden = False
    style.quick_style = True
    _configure_style(style, role_name, ir)
    return style


def _clear_direct_brand_formatting(paragraph) -> None:
    """Remove só família, tamanho e cor diretos; ênfases e links permanecem."""
    for run_properties in paragraph._p.xpath(".//w:r/w:rPr"):
        for tag in _FONT_XML_TAGS:
            for element in list(run_properties.findall(qn(tag))):
                run_properties.remove(element)


def _resolve_logo(ir: BrandIR, asset_root: Path | None) -> Path | None:
    if not ir.assets:
        return None
    token = "logo.primary" if "logo.primary" in ir.assets else sorted(ir.assets)[0]
    candidate = Path(ir.assets[token].path)
    if not candidate.is_absolute() and asset_root is not None:
        candidate = asset_root / candidate
    try:
        resolved = candidate.resolve(strict=True)
    except OSError as error:
        raise DocxBrandError("O logo da marca não foi encontrado para o cabeçalho.") from error
    if not resolved.is_file():
        raise DocxBrandError("O logo da marca precisa ser um arquivo regular.")
    return resolved


def _header_has_image(header) -> bool:
    return bool(header._element.xpath(".//w:drawing") or header._element.xpath(".//w:pict"))


def _office_logo(source: Path, directory: Path) -> tuple[Path, Path | None]:
    """Converte SVG já sanitizado para PNG temporário aceito pelo Word."""
    if source.suffix.casefold() != ".svg":
        return source, None
    import fitz

    descriptor, raw_target = tempfile.mkstemp(prefix=".molda-logo-", suffix=".png", dir=directory)
    os.close(descriptor)
    target = Path(raw_target)
    try:
        with fitz.open(stream=source.read_bytes(), filetype="svg") as document:
            page = document.load_page(0)
            longest = max(page.rect.width, page.rect.height, 1)
            scale = min(4.0, max(1.0, 1024.0 / longest))
            page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=True).save(target)
    except (OSError, RuntimeError, ValueError) as error:
        target.unlink(missing_ok=True)
        raise DocxBrandError("O logo SVG não pôde ser preparado para o Word.") from error
    return target, target


def analyze_docx_brand(path: Path, ir: BrandIR) -> DocxBrandPlan:
    """Valida um DOCX e explica as mudanças de marca antes de aplicá-las."""
    path = path.resolve(strict=True)
    if path.suffix.casefold() != ".docx":
        raise DocxBrandError("Envie um arquivo .docx sem macros.")
    blocking = [item for item in validate_ooxml(path) if item.blocking]
    if blocking:
        raise DocxBrandError(blocking[0].message)
    try:
        document = Document(path)
    except (OSError, ValueError) as error:
        raise DocxBrandError("Não foi possível interpretar o documento Word.") from error

    assignments = _assignments(document, ir)
    heading_name, _ = _semantic_role(ir, "heading")
    body_name, _ = _semantic_role(ir, "body")
    heading_count = sum(item.role_name == heading_name for item in assignments)
    body_count = sum(item.role_name == body_name for item in assignments)
    operations = [
        DocxBrandOperation(
            id="op-001",
            kind="document-styles",
            label_pt="Criar estilos da marca para novos títulos e textos",
            affected_count=2,
        ),
        DocxBrandOperation(
            id="op-002",
            kind="paragraph-styles",
            label_pt=(
                f"Aplicar hierarquia da marca a {heading_count} título(s) e "
                f"{body_count} parágrafo(s)"
            ),
            affected_count=len(assignments),
        ),
        DocxBrandOperation(
            id="op-003",
            kind="page-layout",
            label_pt="Uniformizar margens e respiros das páginas",
            affected_count=len(document.sections),
        ),
    ]
    if document.tables:
        operations.append(
            DocxBrandOperation(
                id="op-004",
                kind="table-styles",
                label_pt="Organizar tabelas com estrutura editável e tipografia da marca",
                affected_count=len(document.tables),
                target_role=body_name,
            )
        )
    if ir.assets:
        operations.append(
            DocxBrandOperation(
                id="op-005",
                kind="header-logo",
                label_pt="Adicionar o logo onde o cabeçalho ainda não possui uma imagem",
                affected_count=len(document.sections),
            )
        )

    warnings = []
    if not assignments:
        warnings.append(
            "O arquivo não possui texto visível; serão aplicados apenas estilos e página."
        )
    if any(paragraph.style is None for paragraph in _all_body_paragraphs(document)):
        warnings.append(
            "Alguns parágrafos não possuem estilo nomeado e usarão o estilo de texto da marca."
        )
    return DocxBrandPlan(
        source=DocxBrandSource(
            filename=path.name,
            sha256=_sha256(path),
            size_bytes=path.stat().st_size,
            paragraph_count=len(assignments),
            table_count=len(document.tables),
            section_count=len(document.sections),
        ),
        brand_revision_id=ir.revision.id,
        operations=operations,
        warnings=warnings,
    )


def apply_docx_brand_plan(
    source: Path,
    out: Path,
    plan: DocxBrandPlan,
    ir: BrandIR,
    *,
    asset_root: Path | None = None,
) -> DocxBrandResult:
    """Aplica o plano em cópia atômica e prova que texto e mídias originais sobreviveram."""
    source = source.resolve(strict=True)
    out = out.resolve()
    if source.suffix.casefold() != ".docx" or out.suffix.casefold() != ".docx":
        raise DocxBrandError("A entrada e a saída precisam usar a extensão .docx.")
    if source == out:
        raise DocxBrandError("O documento original nunca pode ser sobrescrito.")
    if plan.source.sha256 != _sha256(source):
        raise DocxBrandError("Os bytes do Word mudaram depois da análise.")
    if plan.brand_revision_id != ir.revision.id:
        raise DocxBrandError("O plano pertence a outra revisão de marca.")
    blocking = [item for item in validate_ooxml(source) if item.blocking]
    if blocking:
        raise DocxBrandError(blocking[0].message)

    document = Document(source)
    original_content_sha = _content_digest(document)
    original_media = _media_hashes(source)
    assignments = _assignments(document, ir)
    heading_name, _ = _semantic_role(ir, "heading")
    body_name, _ = _semantic_role(ir, "body")
    custom_heading = _ensure_custom_style(document, _CUSTOM_HEADING_STYLE, heading_name, ir)
    custom_body = _ensure_custom_style(document, _CUSTOM_BODY_STYLE, body_name, ir)

    configured: set[tuple[str, str]] = set()
    for assignment in assignments:
        paragraph = assignment.paragraph
        if assignment.force_custom_style:
            paragraph.style = custom_heading
        elif paragraph.style is None:
            paragraph.style = custom_body
        else:
            key = (paragraph.style.style_id, assignment.role_name)
            if key not in configured:
                _configure_style(paragraph.style, assignment.role_name, ir)
                configured.add(key)
        _clear_direct_brand_formatting(paragraph)

    for section in document.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)

    for table in document.tables:
        try:
            table.style = "Table Grid"
        except KeyError:  # pragma: no cover - documentos Word usuais contêm o estilo latente
            pass

    out.parent.mkdir(parents=True, exist_ok=True)
    logo = _resolve_logo(ir, asset_root)
    office_logo: Path | None = None
    temporary_logo: Path | None = None
    if logo is not None:
        office_logo, temporary_logo = _office_logo(logo, out.parent)
        seen_headers: set[str] = set()
        for section in document.sections:
            header = section.header
            part_name = str(header.part.partname)
            if part_name in seen_headers:
                continue
            seen_headers.add(part_name)
            if _header_has_image(header):
                continue
            paragraph = header.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            try:
                paragraph.add_run().add_picture(str(office_logo), width=Inches(1.15))
            except (OSError, ValueError) as error:
                if temporary_logo is not None:
                    temporary_logo.unlink(missing_ok=True)
                raise DocxBrandError(
                    "O logo não pôde ser inserido no cabeçalho do Word."
                ) from error

    descriptor, raw_temp = tempfile.mkstemp(
        prefix=f".{out.stem}-",
        suffix=".docx",
        dir=out.parent,
    )
    os.close(descriptor)
    temporary = Path(raw_temp)
    try:
        document.save(temporary)
        blocking = [item for item in validate_ooxml(temporary) if item.blocking]
        if blocking:
            raise DocxBrandError(blocking[0].message)
        branded_document = Document(temporary)
        branded_content_sha = _content_digest(branded_document)
        if branded_content_sha != original_content_sha:
            raise DocxBrandError("A aplicação foi cancelada porque o conteúdo textual mudou.")
        if not original_media.issubset(_media_hashes(temporary)):
            raise DocxBrandError("A aplicação foi cancelada porque uma imagem original se perdeu.")
        branded_sha = _sha256(temporary)
        os.replace(temporary, out)
    finally:
        temporary.unlink(missing_ok=True)
        if temporary_logo is not None:
            temporary_logo.unlink(missing_ok=True)

    return DocxBrandResult(
        source_sha256=plan.source.sha256,
        branded_sha256=branded_sha,
        output_filename=out.name,
        applied_operation_ids=[operation.id for operation in plan.operations],
        content_preserved=True,
        content_sha256=original_content_sha,
    )
