from io import BytesIO

from docx import Document

from brand_api.models import Job
from brand_api.worker import run_next_job
from brand_runtime import validate_ooxml


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Proposta para o novo projeto")
    document.add_paragraph("Este conteúdo deve continuar editável depois de receber a marca.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Fase"
    table.cell(0, 1).text = "Prazo"
    table.cell(1, 0).text = "Descoberta"
    table.cell(1, 1).text = "5 dias"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _run_one(client):
    state = client.app.state
    return run_next_job(
        state.session_factory,
        storage=state.storage,
        exporter=state.exporter,
        settings=state.settings,
    )


def test_docx_branding_api_worker_preserva_e_publica(client, compiled, tmp_path):
    source = _docx_bytes()
    response = client.post(
        f"/v1/brand-revisions/{compiled['brandRevisionId']}/docx-brandings",
        files={"file": ("proposta do time.docx", source, "application/octet-stream")},
    )
    assert response.status_code == 202, response.text
    analysis_job_id = response.json()["jobId"]
    with client.app.state.session_factory() as session:
        queued = session.get(Job, analysis_job_id)
        assert queued is not None and queued.kind == "docx-brand-analyze"
        source_sha256 = queued.params["sourceSha256"]
    assert client.app.state.storage.get(source_sha256) == source

    assert _run_one(client)
    analysis = client.get(f"/v1/jobs/{analysis_job_id}").json()
    assert analysis["status"] == "succeeded", analysis
    plan = analysis["result"]["plan"]
    assert plan["source"]["filename"] == "proposta do time.docx"
    assert plan["source"]["tableCount"] == 1
    assert {operation["kind"] for operation in plan["operations"]} >= {
        "document-styles",
        "paragraph-styles",
        "page-layout",
        "header-logo",
    }

    response = client.post(f"/v1/jobs/{analysis_job_id}/docx-brandings")
    assert response.status_code == 202, response.text
    apply_job_id = response.json()["jobId"]
    assert _run_one(client)
    applied = client.get(f"/v1/jobs/{apply_job_id}").json()
    assert applied["status"] == "succeeded", applied
    result = applied["result"]
    assert result["kind"] == "docx-brand-apply"
    assert result["filename"] == "proposta do time-com-marca.docx"
    assert result["brandResult"]["contentPreserved"] is True

    download = client.get(result["url"])
    assert download.status_code == 200
    assert download.headers["content-type"].endswith("wordprocessingml.document")
    output = tmp_path / "branded.docx"
    output.write_bytes(download.content)
    assert not [item for item in validate_ooxml(output) if item.blocking]
    document = Document(output)
    assert document.paragraphs[0].text == "Proposta para o novo projeto"
    assert document.paragraphs[0].style.name == "Molda Título"
    assert document.tables[0].cell(1, 0).text == "Descoberta"
    assert document.sections[0].header._element.xpath(".//w:drawing")
    assert client.app.state.storage.get(source_sha256) == source
    assert not list(client.app.state.settings.work_dir.iterdir())


def test_docx_branding_recusa_arquivo_estado_e_auth(client, anon, compiled):
    endpoint = f"/v1/brand-revisions/{compiled['brandRevisionId']}/docx-brandings"
    invalid = client.post(
        endpoint,
        files={"file": ("falso.docx", b"not-a-zip", "application/octet-stream")},
    )
    assert invalid.status_code == 400
    assert invalid.json()["detail"] == "Envie um documento Word .docx válido e sem macros."

    unauthorized = anon.post(
        endpoint,
        files={"file": ("proposta.docx", _docx_bytes(), "application/octet-stream")},
    )
    assert unauthorized.status_code == 401

    missing = client.post("/v1/jobs/job_000000000000/docx-brandings")
    assert missing.status_code == 404
